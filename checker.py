from mistralai import Mistral
from mistral.exceptions import MistralException
import requests
import docx
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
import time

doc_path = "УПД _не трогать.docx"

# Ваши API ключи прямо в коде
API_KEYS = [
#Добавьте сюда несколько api ключей для ротации
]

current_key_index = 0
key_usage_count = {}  # Счетчик использования каждого ключа

def get_next_api_key():
    """Получает следующий API ключ (ротация по кругу)"""
    global current_key_index
    
    if not API_KEYS:
        raise ValueError("Нет доступных API ключей!")
    
    if current_key_index >= len(API_KEYS):
        current_key_index = 0
    
    key = API_KEYS[current_key_index]
    
    # Увеличиваем счетчик использования
    if key not in key_usage_count:
        key_usage_count[key] = 0
    key_usage_count[key] += 1
    
    print(f"Используем ключ №{current_key_index + 1} (индекс {current_key_index})")
    print(f"Этот ключ использовался {key_usage_count[key]} раз")
    
    # Переключаемся на следующий ключ для следующего вызова
    current_key_index = (current_key_index + 1) % len(API_KEYS)
    
    return key

def create_mistral_client():
    """Создает клиент Mistral с текущим API ключом"""
    api_key = get_next_api_key()
    return Mistral(api_key=api_key), api_key

def safe_mistral_call(func, *args, max_retries=None, **kwargs):
    """
    Безопасный вызов функции Mistral с автоматической ротацией ключей при ошибках
    """
    global current_key_index
    
    if max_retries is None:
        max_retries = len(API_KEYS) * 2
    
    original_key_index = current_key_index
    
    for attempt in range(max_retries):
        try:
            print(f"\n--- Попытка {attempt + 1}/{max_retries} ---")
            result = func(*args, **kwargs)
            print(f"Успешно использован ключ {current_key_index}")
            return result
            
        except MistralException as e:
            error_msg = str(e).lower()
            print(f"Ошибка Mistral: {e}")
            
            # Проверяем типичные ошибки лимитов
            is_limit_error = any(word in error_msg for word in [
                "quota", "limit", "exceeded", "unauthorized", 
                "invalid", "rate", "429", "too many"
            ])
            
            if is_limit_error:
                print(f"Ключ {current_key_index} исчерпал лимит или недействителен.")
                
                # Ждем перед следующей попыткой
                wait_time = 2 + attempt  # Увеличиваем время ожидания с каждой попыткой
                print(f"Ждем {wait_time} секунд перед следующей попыткой...")
                time.sleep(wait_time)
                
                # Продолжаем со следующим ключом
                continue
            else:
                # Другие ошибки
                print(f"Другая ошибка API: {e}")
                raise
                
        except Exception as e:
            print(f"Неожиданная ошибка: {e}")
            raise
    
    # Если все попытки исчерпаны
    used_keys_str = "\n".join([f"Ключ {i+1}: {key[:10]}... использован {count} раз" 
                               for i, (key, count) in enumerate(key_usage_count.items())])
    print(f"\nВсе ключи исчерпали лимиты. Статистика использования:")
    print(used_keys_str)
    raise Exception(f"Все {max_retries} попыток исчерпаны. Проверьте API ключи.")

def set_format(doc_path):
    doc = Document(doc_path)
    #междустрочный интервал
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
    
    for section in doc.sections:
        section.left_margin = Cm(3) #30мм
        section.right_margin = Cm(1.5) # 15мм
        section.top_margin = Cm(2) # 20мм
        section.bottom_margin = Cm(2) # 20мм
        
    doc.save(doc_path)

def ai_log_admin(doc_path):
    """Первая проверка ИИ"""
    
    #for docks
    gost = "ОС ТУСУР 01-2021"
    doc = Document(doc_path)
    text = []
          
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():  
            text.append(f"{i}: {paragraph.text}")
    
    text = '\n'.join(text)
   
    system_prompt = """
    Ты — эксперт по технической документации и стандартам ГОСТ. Проанализируй 
    предоставленный текст документа на соответствие требованиям ГОСТ по следующим аспектам:
    
    КОНТЕКСТНЫЕ И СМЫСЛОВЫЕ ПРОВЕРКИ:
     1. Семантические ошибки:
        - Проверь точность технических терминов и их соответствие контексту
        - Выяви некорректные формулировки, искажающие смысл требований
        - Найди двусмысленные трактовки, допускающие multiple интерпретации
    2. Логические противоречия:
        - Обнаружь взаимоисключающие требования в разных частях документа
        - Проверь логическую последовательность предписаний
        - Выяви противоречия между числовыми значениями и их описаниями
    3. Стилистические нормы:
        - Оцени соответствие официально-деловому стилю ГОСТ
        - Проверь единообразие терминологии по всему документу
        - Выяви нарушения в структуре формулировок
    4. Контекстные зависимости:
        - Проверь корректность ссылок на другие разделы/стандарты
        - Выяви требования, зависящие от контекста, но не имеющие четких условий
        - Обнаружь несоответствия между описательными и нормативными частями
    5. Орфографические ошибки (неправильное написание слов)
        
    Формат твоего ответа: [Номер параграфа][Ошибка(приведенные выше)] - Исправленный текст
    
    - Исправляй, сохраняя исходный смысл
    - Если сомневаешься - не исправляй
    - Фокусируйся только на существенных нарушениях
    
    ПИШИ ТОЛЬКО В ТАКОМ ФОРМАТЕ: [Номер параграфа][Ошибка] - Исправленный текст
    не пиши причины исправления
    не пиши старый текст, а только исправленный
    не исправляй фио студентов и преподавателей!!!
    """
    
    try:
        client, current_key = create_mistral_client()
        
        def call_mistral():
            return client.chat.complete(
                model="mistral-large-latest",
                messages=[
                    {"role": "system",
                    "content": f"ГОСТ: {gost}\n {system_prompt}"},
                    {"role": "user",
                    "content":f"Текст для проверки:\n{text}"}
                ],
                temperature=0.3,
                max_tokens=3000
            )
        
        # Используем безопасный вызов с ротацией ключей
        chat_response = safe_mistral_call(call_mistral)
        
        result = chat_response.choices[0].message.content
        log = [result]
        print(f"ai_log_admin успешно выполнен, получено {len(result)} символов")
        return log
    
    except Exception as e:
        print(f"Ошибка в ai_log_admin: {e}")
        return None            

def checker_AIs(doc_path):
    """Основная функция проверки с двумя ИИ"""
    
    #for docks
    gost = "ОС ТУСУР 01-2021"
    doc = Document(doc_path)
    text = []
          
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():  
            text.append(f"{i}: {paragraph.text}")
    
    text = '\n'.join(text)   
    
    sys_prompt = """
    Ты являешься независимым экспертом по оцениванию ГОСТов
    Твоя задача брать 2 решения ии, которые являются такими же экспертами и сравнивать их 
    ответы со своим и выдавать оптимальное решение по этим аспектам ГОСТов:
    
    КОНТЕКСТНЫЕ И СМЫСЛОВЫЕ ПРОВЕРКИ:
     1. Семантические ошибки:
        - Проверь точность технических терминов и их соответствие контексту
        - Выяви некорректные формулировки, искажающие смысл требований
        - Найди двусмысленные трактовки, допускающие multiple интерпретации
    2. Логические противоречия:
        - Обнаружь взаимоисключающие требования в разных частях документа
        - Проверь логическую последовательность предписаний
        - Выяви противоречия между числовыми значениями и их описаниями
    3. Стилистические нормы:
        - Оцени соответствие официально-деловому стилю ГОСТ
        - Проверь единообразие терминологии по всему документу
        - Выяви нарушения в структуре формулировок
    4. Контекстные зависимости:
        - Проверь корректность ссылок на другие разделы/стандарты
        - Выяви требования, зависящие от контекста, но не имеющие четких условий
        - Обнаружь несоответствия между описательными и нормативные части
    5. Орфографические ошибки (неправильное написание слов)
        
    Формат твоего ответа: [Номер параграфа]Исправленный текст
    
    - Исправляй, сохраняя исходный смысл
    - Если сомневаешься - не исправляй
    - Фокусируйся только на существенных нарушениях
    
    ПИШИ ТОЛЬКО В ТАКОМ ФОРМАТЕ: [Номер параграфа]Исправленный текст
    не пиши причины исправления
    не пиши старый текст, а только исправленный
    не подчеркивай исправленную часть в параграфе 
    не исправляй фио студентов и преподавателей!!!
    """
    
    print("\n" + "="*50)
    print("Запуск первой проверки ИИ...")
    log1 = ai_log_admin(doc_path)
    
    print("\n" + "="*50)
    print("Запуск второй проверки ИИ...")
    log2 = ai_log_admin(doc_path)
    
    if not log1 or not log2:
        print("Одна из проверок ИИ не удалась")
        return None
    
    try:
        client, current_key = create_mistral_client()
        
        def call_mistral():
            return client.chat.complete(
                model="mistral-large-latest",
                messages=[
                    {"role": "system",
                    "content": f"ГОСТ: {gost}\n {sys_prompt} \nОтвет 1го ИИ: {log1}\nОтвет 2го ИИ: {log2}"},
                    {"role": "user",
                    "content":f"Текст для проверки:\n{text}"}
                ],
                temperature=0.3,
                max_tokens=3000
            )
        
        # Используем безопасный вызов с ротацией ключей
        print("\n" + "="*50)
        print("Запуск финальной проверки ИИ...")
        chat_response = safe_mistral_call(call_mistral)
        
        result = chat_response.choices[0].message.content
        print(f"checker_AIs успешно выполнен, получено {len(result)} символов")
        return result
        
    except Exception as e:
        print(f"Ошибка в checker_AIs: {e}")
        return None    

def parsing_ai_response(otvet):
    """Парсинг ответа ИИ"""
    paragraphs_dict = {}
    
    lines = otvet.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if line.startswith('[') and ']' in line:
            # Разделяем на номер и текст
            bracket_pos = line.find(']')
            number_part = line[1:bracket_pos]  # Текст между [ и ]
            
            if number_part.isdigit():
                paragraph_number = int(number_part)
                paragraph_text = line[bracket_pos + 1:].strip()
                
                # Добавляем в словарь
                paragraphs_dict[paragraph_number] = paragraph_text
    
    return paragraphs_dict
            
def change_context(directory, doc_path):
    """Изменение документа на основе исправлений"""
    doc = Document(doc_path)
    
    changes_count = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if i in directory:
            paragraph.text = directory[i]
            changes_count += 1
            print(f"Исправлен параграф {i}")
    
    doc.save(doc_path)
    print(f"Всего исправлено {changes_count} параграфов")

if __name__ == "__main__":
    print("="*50)
    print("=== Запуск проверки ГОСТов ===")
    print(f"Доступно {len(API_KEYS)} API ключей")
    print("="*50)
    

    
    
    
    otvet = checker_AIs(doc_path)
    
    if otvet:
        print("\n" + "="*50)
        print("Обработка результатов...")
        
        directory_changes = parsing_ai_response(otvet=otvet)
        print(f"Найдено {len(directory_changes)} исправлений")
        
        change_context(directory=directory_changes, doc_path=doc_path)
        
        print("\n" + "="*50)
        print("Результат проверки:")
        print("-"*50)
        print(otvet[:500] + "..." if len(otvet) > 500 else otvet)
        print("-"*50)
        
        set_format(doc_path)
        print("Форматирование документа завершено!")
        print("\n" + "="*50)
        print("Проверка завершена успешно!")
        
        # Выводим статистику использования ключей
        print("\nСтатистика использования ключей:")
        total_uses = sum(key_usage_count.values())
        print(f"Всего использований ключей: {total_uses}")
        for i, key in enumerate(API_KEYS):
            count = key_usage_count.get(key, 0)
            print(f"Ключ {i+1}: использован {count} раз")
    else:
        print("Не удалось получить ответ от ИИ. Проверьте:")
        print("1. Доступность интернета")
        print("2. Активность API ключей")
        print("3. Правильность формата документа")